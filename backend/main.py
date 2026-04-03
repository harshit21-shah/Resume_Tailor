from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv
import os
import json
import re
import uuid
import docx
from docx2pdf import convert
import pdfplumber
from pdf2docx import Converter
import asyncio
import base64

load_dotenv()

app = FastAPI(title="Resume Tailor API")
print("BACKEND VERSION 2.0 - DL SYSTEM ACTIVE")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class TailorRequest(BaseModel):
    resume: str
    job_description: str


class Change(BaseModel):
    section: str
    description: str


class Replacement(BaseModel):
    original: str
    new: str

class DownloadRequest(BaseModel):
    file_id: str
    replacements: list[Replacement]


class TailorResponse(BaseModel):
    tailored_resume: str
    changes: list[Change]
    match_score: int
    matched_keywords: list[str]
    missing_keywords: list[str]
    replacements: list[Replacement]


SYSTEM_PROMPT = """You are an expert ATS resume optimizer and career coach with deep knowledge of how applicant tracking systems parse resumes.

Your job is to surgically tailor a candidate's resume to match a specific job description — improving ATS score and keyword alignment without fabricating any experience.

STRICT RULES:
1. NEVER change: candidate name, contact info, dates (including 'Present' or specific years), company names, job titles, university names, or degree names.
2. LINE STRUCTURE: You must preserve the exact line-by-line structure. NEVER merge a company line with a job title line or a date line.
3. DO modify:
   - Profile Summary: Completely rewrite to mirror the JD's language, tone, and key requirements.
   - Skills section: Reorder and keyword-match skills to JD. 
   - Experience/Project bullets: Swap synonyms to match JD keywords while keeping factual metrics.
4. Keep identical document structure, whitespace, and formatting.
5. Provide exact string replacements for modified text only.
6. The tailored resume must be honest and 100% defensible.
7. NEUTER CHARACTERS: Use standard dashes (-) and pipes (|) as found in the input. Don't convert them randomly.
"""

USER_PROMPT_TEMPLATE = """RESUME:
{resume}

JOB DESCRIPTION:
{job_description}

Analyze the JD carefully. Identify the top keywords, required skills, and tone. Then tailor the resume.
Also provide a strict list of exact substring replacements to map your changes from the original text. For each modified paragraph or bullet point, give the EXACT original text as a string and the EXACT new text.

Respond with ONLY a valid JSON object in this exact structure (no markdown fences, no explanation before or after):
{{
  "tailored_resume": "the complete tailored resume as a string with \\n for newlines",
  "changes": [
    {{"section": "Profile Summary", "description": "Specific description of what changed and why it matches the JD"}},
    {{"section": "Skills", "description": "Specific description of what was reordered or added"}},
    {{"section": "Actorius bullet 2", "description": "Specific description of phrasing change"}}
  ],
  "replacements": [
    {{"original": "exact original paragraph text (no extra spaces)", "new": "exact new paragraph text"}}
  ],
  "match_score": 78,
  "matched_keywords": ["keyword1", "keyword2", "keyword3"],
  "missing_keywords": ["keyword4", "keyword5"]
}}

match_score should be an integer 0-100 representing ATS alignment after tailoring.
matched_keywords: top 8 JD keywords present in the tailored resume.
missing_keywords: up to 5 important JD keywords the candidate genuinely cannot claim."""


@app.get("/health")
def health():
    return {"status": "ok"}

def extract_text_from_docx(doc) -> str:
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)
    return "\n".join(lines)

def apply_replacements_to_docx(doc, replacements: list[Replacement]):
    def replace_in_para(paragraph):
        doc_text = paragraph.text
        for rep in replacements:
            old_t = rep.original.strip()
            new_t = rep.new.strip()
            if not old_t or not new_t or old_t == new_t:
                continue
            
            # Normalize common dash variations for matching
            normalized_doc_text = doc_text.replace('–', '-').replace('—', '-')
            normalized_old_t = old_t.replace('–', '-').replace('—', '-')

            if normalized_old_t in normalized_doc_text:
                replaced = False
                for run in paragraph.runs:
                    norm_run_text = run.text.replace('–', '-').replace('—', '-')
                    if normalized_old_t in norm_run_text:
                        run.text = run.text.replace(old_t, new_t)
                        replaced = True
                        doc_text = paragraph.text 
                
                if not replaced:
                    first_run_style = None
                    if paragraph.runs:
                        first_run_style = paragraph.runs[0]
                    
                    if normalized_old_t == normalized_doc_text:
                        paragraph.text = new_t
                    else:
                        paragraph.text = paragraph.text.replace(old_t, new_t)
                    
                    # Restore formatting on the new run(s)
                    for run in paragraph.runs:
                        if first_run_style:
                            run.bold = first_run_style.bold
                            run.italic = first_run_style.italic
                            run.font.name = first_run_style.font.name
                            run.font.size = first_run_style.font.size
                    replaced = True
                
                if replaced:
                    break

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

async def delete_after_delay(path: str, delay: int = 3600):
    await asyncio.sleep(delay)
    cleanup_file(path)

@app.post("/upload")
async def upload_resume(file: UploadFile = File(...), background_tasks: BackgroundTasks = BackgroundTasks()):
    if not file.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    file_id = uuid.uuid4().hex[:8]
    ext = os.path.splitext(file.filename)[1].lower()
    
    raw_path = f"temp_raw_{file_id}{ext}"
    with open(raw_path, "wb") as f:
        f.write(await file.read())
        
    extracted_text = ""
    target_docx = f"temp_{file_id}.docx"
    
    try:
        if ext == '.pdf':
            with pdfplumber.open(raw_path) as pdf:
                extracted_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            cv = Converter(raw_path)
            cv.convert(target_docx, start=0, end=None)
            cv.close()
        else:
            doc = docx.Document(raw_path)
            extracted_text = extract_text_from_docx(doc)
            doc.save(target_docx)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")
    finally:
        if os.path.exists(raw_path):
            os.remove(raw_path)
            
    background_tasks.add_task(delete_after_delay, target_docx, 3600)
        
    return {"file_id": file_id, "text": extracted_text}

@app.post("/ocr-jd")
async def ocr_job_description(file: UploadFile = File(...)):
    print(f"DEBUG: Starting OCR for file: {file.filename}")
    if not file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.webp')):
        raise HTTPException(status_code=400, detail="Only images are supported for JD OCR")
    
    contents = await file.read()
    print(f"DEBUG: File read complete. Size: {len(contents)} bytes")
    base64_image = base64.b64encode(contents).decode('utf-8')
    print("DEBUG: Base64 encoding complete. Sending to Groq...")
    
    try:
        completion = client.chat.completions.create(
            model="llama-3.2-11b-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all the text regarding the job description from this image accurately. Do not include any other text or explanations. Return only the extracted text."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            timeout=30.0 # Add a 30s timeout
        )
        extracted_text = completion.choices[0].message.content.strip()
        print(f"DEBUG: Groq Vision success. Extracted {len(extracted_text)} chars.")
        return {"text": extracted_text}
    except Exception as e:
        print(f"DEBUG: OCR Exception: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to OCR image: {str(e)}")

def cleanup_file(path: str):
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


@app.post("/tailor", response_model=TailorResponse)
def tailor_resume(req: TailorRequest):
    if not req.resume.strip():
        raise HTTPException(status_code=400, detail="Resume cannot be empty")
    if not req.job_description.strip():
        raise HTTPException(status_code=400, detail="Job description cannot be empty")
    if len(req.resume) > 8000:
        raise HTTPException(status_code=400, detail="Resume too long (max 8000 chars)")
    if len(req.job_description) > 5000:
        raise HTTPException(status_code=400, detail="Job description too long (max 5000 chars)")

    prompt = USER_PROMPT_TEMPLATE.format(
        resume=req.resume,
        job_description=req.job_description
    )

    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        max_tokens=4096,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
    )

    raw = completion.choices[0].message.content.strip()

    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    raw = raw.strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', raw)
        if match:
            parsed = json.loads(match.group(0))
        else:
            raise HTTPException(status_code=500, detail="Failed to parse AI response")

    return TailorResponse(
        tailored_resume=parsed.get("tailored_resume", ""),
        changes=[Change(**c) for c in parsed.get("changes", [])],
        match_score=parsed.get("match_score", 0),
        matched_keywords=parsed.get("matched_keywords", []),
        missing_keywords=parsed.get("missing_keywords", []),
        replacements=[Replacement(**r) for r in parsed.get("replacements", [])]
    )

@app.post("/generate-docx")
def generate_docx(req: DownloadRequest, background_tasks: BackgroundTasks):
    template_path = f"temp_{req.file_id}.docx"
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Original resume template not found.")
        
    doc = docx.Document(template_path)
    apply_replacements_to_docx(doc, req.replacements)
                
    export_id = uuid.uuid4().hex
    output_filename = f"export_{export_id}.docx"
    doc.save(output_filename)
    
    background_tasks.add_task(delete_after_delay, output_filename, 3600)
    return {"url": f"/dl/{export_id}.docx"}

@app.post("/generate-pdf")
def generate_pdf(req: DownloadRequest, background_tasks: BackgroundTasks):
    template_path = f"temp_{req.file_id}.docx"
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Original resume template not found.")
        
    doc = docx.Document(template_path)
    apply_replacements_to_docx(doc, req.replacements)
                
    export_id = uuid.uuid4().hex
    temp_docx = f"temp_export_{export_id}.docx"
    temp_pdf = f"export_{export_id}.pdf"
    
    doc.save(temp_docx)
    try:
        convert(temp_docx, temp_pdf)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF Conversion failed: {str(e)}")
    finally:
        cleanup_file(temp_docx)
        
    background_tasks.add_task(delete_after_delay, temp_pdf, 3600)
    return {"url": f"/dl/{export_id}.pdf"}

@app.get("/dl/{file_key}")
def get_export(file_key: str):
    # file_key will be like "UUID.pdf"
    export_id = file_key.split('.')[0]
    ext = ".pdf" if file_key.lower().endswith(".pdf") else ".docx"
    actual_file = f"export_{export_id}{ext}"
    
    if not os.path.exists(actual_file):
        raise HTTPException(status_code=404, detail="File not found or expired")
    
    media_type = "application/pdf" if ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    pretty_name = "Tailored_Resume.pdf" if ext == ".pdf" else "Tailored_Resume.docx"
    
    from fastapi.responses import FileResponse
    response = FileResponse(actual_file, media_type=media_type)
    response.headers["Content-Disposition"] = f'attachment; filename="{pretty_name}"'
    return response
